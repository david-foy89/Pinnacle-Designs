"""Convert contract Markdown files to Word (.docx) documents."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\[[^\]]+\])"
)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_inline_runs(paragraph, text: str) -> None:
    if not text:
        return

    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    return bool(re.fullmatch(r"[\s:\-|]+", stripped))


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if is_table_separator(line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            value = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            add_inline_runs(paragraph, value)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True


def add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(level=level)
    add_inline_runs(heading, text)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_defaults(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        if is_table_row(stripped):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table_rows(table_lines))
            continue

        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            for item in items:
                paragraph = doc.add_paragraph(style="List Bullet")
                add_inline_runs(paragraph, item)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            for item in items:
                paragraph = doc.add_paragraph(style="List Number")
                add_inline_runs(paragraph, item)
            continue

        paragraph_lines = [stripped.rstrip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if (
                not nxt_stripped
                or nxt_stripped == "---"
                or nxt_stripped.startswith("#")
                or is_table_row(nxt_stripped)
                or re.match(r"^[-*]\s+", nxt_stripped)
                or re.match(r"^\d+\.\s+", nxt_stripped)
            ):
                break
            paragraph_lines.append(nxt_stripped.rstrip())
            i += 1

        paragraph = doc.add_paragraph()
        for idx, part in enumerate(paragraph_lines):
            if idx:
                paragraph.add_run("\n")
            add_inline_runs(paragraph, part)

    doc.save(docx_path)


def main() -> int:
    contracts_dir = Path(__file__).resolve().parents[1] / "contracts"
    targets = sorted(
        p for p in contracts_dir.glob("*.md") if p.name.lower() != "readme.md"
    )

    if not targets:
        print("No contract markdown files found.", file=sys.stderr)
        return 1

    for md_path in targets:
        docx_path = md_path.with_suffix(".docx")
        convert_markdown_to_docx(md_path, docx_path)
        print(f"Created {docx_path.name}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
